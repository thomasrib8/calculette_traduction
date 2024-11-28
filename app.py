from flask import Flask, render_template, request, redirect, url_for, session
from werkzeug.utils import secure_filename
import os
from datetime import timedelta
import docx
from flask_wtf import FlaskForm
from wtforms import StringField, PasswordField, SubmitField
from wtforms.validators import DataRequired

app = Flask(__name__)
app.config['UPLOAD_FOLDER'] = 'uploads'
app.config['SECRET_KEY'] = 'your_secret_key'
app.config['PERMANENT_SESSION_LIFETIME'] = timedelta(minutes=30)

# Credentials for login
USERNAME = "admin"  # Replace with your desired username
PASSWORD = "Roue2021*"  # Replace with your desired password

# Ensure the upload folder exists
os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)


class LoginForm(FlaskForm):
    username = StringField("Username", validators=[DataRequired()])
    password = PasswordField("Password", validators=[DataRequired()])
    submit = SubmitField("Login")


def get_docx_stats(file_path):
    doc = docx.Document(file_path)
    words = 0
    characters = 0
    paragraphs = len(doc.paragraphs)
    for paragraph in doc.paragraphs:
        words += len(paragraph.text.split())
        characters += len(paragraph.text)
    pages = len(doc.sections)
    return words, characters, pages, paragraphs


def calculate_translation_time(words, paragraphs, group_size):
    step1_time = words * 0.00156
    group_time_map = {
        1: 1.5, 2: 1.8, 3: 2.5, 4: 3.0, 5: 3.2,
        6: 3.3, 7: 3.5, 8: 4.0, 9: 3.8, 10: 4.0
    }
    step2_time = paragraphs * group_time_map.get(group_size, 3.0)
    total_time_sec = step1_time + step2_time
    return timedelta(seconds=total_time_sec), total_time_sec


def calculate_translation_cost(words, characters, translation_time_min):
    tokens = words * 2
    step1_cost = tokens * 0.0000015
    step2_cost = characters * 0.000021
    step3_cost = translation_time_min * 0.005161
    return step1_cost + step2_cost + step3_cost


def calculate_review_cost(words, reviewer_choice):
    """
    Calculate review cost based on the reviewer choice.
    - TOBY or TOBY+MIKE: $0.0069 per word
    - MIKE: $0 (no review cost)
    """
    if reviewer_choice in ["TOBY", "TOBY+MIKE"]:
        return words * 0.024
    elif reviewer_choice == "MIKE":
        return 0
    else:
        raise ValueError("Invalid reviewer choice")


@app.route('/', methods=['GET', 'POST'])
def login():
    """Login route to authenticate user."""
    if 'logged_in' in session and session['logged_in']:
        return redirect(url_for('index'))

    form = LoginForm()
    if form.validate_on_submit():
        if form.username.data == USERNAME and form.password.data == PASSWORD:
            session['logged_in'] = True
            return redirect(url_for('index'))
        else:
            return render_template('login.html', form=form, error="Invalid username or password.")

    return render_template('login.html', form=form)


@app.route('/logout')
def logout():
    """Logout route to redirect user to login page."""
    session.pop('logged_in', None)
    return redirect(url_for('login'))


@app.route('/home', methods=['GET', 'POST'])
def index():
    """Main page after login."""
    if 'logged_in' not in session or not session['logged_in']:
        return redirect(url_for('login'))

    if request.method == 'POST':
        file = request.files['file']
        group_size = int(request.form['group_size'])
        reviewer_choice = request.form['reviewer']

        # Save the uploaded file
        file_path = os.path.join(app.config['UPLOAD_FOLDER'], secure_filename(file.filename))
        file.save(file_path)

        # Process the document
        words, characters, pages, paragraphs = get_docx_stats(file_path)
        translation_time, translation_time_sec = calculate_translation_time(words, paragraphs, group_size)
        translation_time_min = translation_time_sec / 60
        translation_cost = calculate_translation_cost(words, characters, translation_time_min)
        review_cost = calculate_review_cost(words, reviewer_choice)
        total_cost = translation_cost + review_cost

        return render_template('result.html', words=words, characters=characters, pages=pages,
                               paragraphs=paragraphs, translation_time=str(translation_time),
                               translation_cost=round(translation_cost, 6), review_cost=round(review_cost, 2),
                               total_cost=round(total_cost, 6))

    return render_template('index.html')


if __name__ == '__main__':
    app.run(debug=True)
